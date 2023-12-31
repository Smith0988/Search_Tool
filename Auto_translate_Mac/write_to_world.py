import re
import sys

import nltk
import os
import pandas as pd
from docx import Document
import requests
from bs4 import BeautifulSoup
from google_translate import translate_with_google_translate


copyright_text = "Bản quyền © 2023 Minghui.org. Mọi quyền được bảo lưu."

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

csv_filename = resource_path("link_eng_vn_gct.csv")
csv_filename_dic = resource_path("dic_eng_vn_data.csv")

def find_vietnamese_link_1(english_link):
    # Đọc dữ liệu từ file CSV
    df = pd.read_csv(csv_filename)
    match = ""
    if re.match(r'https?://', english_link):
        match = re.search(r'(en\..*?\.html)', english_link)
    if match:
        english_link_1 = match.group(1)
    else:
        english_link_1 = english_link
    # Tìm link tiếng Anh trong cột 1 sử dụng biểu thức chính quy
    pattern = re.escape(english_link_1)  # Chuyển `english_link` thành biểu thức chính quy
    matches = df[df.iloc[:, 0].str.contains(pattern, regex=True, case=False)]

    if not matches.empty:
        vietnamese_link = matches.iloc[0, 1]
        if not vietnamese_link.startswith("http://") and not vietnamese_link.startswith("https://"):
           vietnamese_link = "https:" + vietnamese_link  # Thêm schema "https:" nếu cần

        return vietnamese_link
    else:
        return "Can not find Vietnamese Link"


def get_related_link(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        article_contents_1 = soup.find('div', class_='article-body-content')
        article_contents_2 = soup.find('div', class_='articleZhengwen geo cBBlack')
        if article_contents_1:
            article_contents = article_contents_1
        else:
            article_contents = article_contents_2
        if article_contents:
            paragraphs = article_contents.find_all(['p', 'h3'])
            valid_paragraphs = []
            start_collecting = False
            related_links = []
            for paragraph in paragraphs:
                text = paragraph.get_text(strip=True)

                if "Related article" in text or "Related Article" in text or "Related Report" in text or "Related report" in text:
                    start_collecting = True

                if start_collecting:
                    if 'splitted' in paragraph.get("class", []):
                        span_sections = paragraph.find_all('span', class_='section')
                        for span in span_sections:
                            text_1 = span.get_text(strip=True)
                            valid_paragraphs.append(text_1)
                            for link in paragraph.find_all('a', href=True):
                                related_links.append(link['href'])
                    else:
                        valid_paragraphs.append(text)
                        for link in paragraph.find_all('a', href=True):
                            related_links.append(link['href'])

            if valid_paragraphs:
                article_content = "\n".join(valid_paragraphs)
                return valid_paragraphs, related_links
            else:
                return [], related_links
        else:
            return [], []
    else:
        return [], []

def find_vietnamese_link(english_link):
    first_text = "Bài liên quan:"
    all_link = []
    result_link = []
    result_link.append(first_text)
    # Đọc dữ liệu từ file CSV
    related_content, related_link = get_related_link(english_link)
    if related_content:
        for link in related_link:
            all_link.append(link)
        for link in all_link:
            result_link.append(find_vietnamese_link_1(link))
        return related_content, result_link, all_link
    else:
        return []
def find_vietnamese_sentence(english_sentence):

    # Đọc dữ liệu từ file CSV
    df = pd.read_csv(csv_filename_dic)

    # Tìm link tiếng Anh trong cột 1
    row = df[df.iloc[:, 0] == english_sentence]

    if not row.empty:
        vietnamese_sentence = row.iloc[0, 1]
        return vietnamese_sentence
    else:
        return []
def get_vn_article_title(url):
    try:
        #kiem tra
        if not url.startswith("http://") and not url.startswith("https://"):
           url = "https:" + url  # Thêm schema "https:" nếu cần
        # Tải nội dung của trang web
        response = requests.get(url)

        # Kiểm tra nếu yêu cầu thành công (status code 200)
        if response.status_code == 200:
            # Sử dụng BeautifulSoup để phân tích cú pháp trang web
            soup = BeautifulSoup(response.content, 'html.parser')

            # Tìm thẻ div có class là 'article-title'
            article_title_tag = soup.find('h1', class_='articleTitle cABlue')

            # Kiểm tra nếu tồn tại thẻ và lấy nội dung text của tiêu đề
            if article_title_tag:
                article_title = article_title_tag.text.strip()
                return article_title
            else:
                return "Can not find Vietnamese Link"
        else:
            return "Can not find Vietnamese Link"
    except Exception as e:
        return "Can not find Vietnamese Link"

def get_en_article_title(url):
    response = requests.get(url)
    article_title = "Can not get title from web"
    if response.status_code == 200:
        # Sử dụng BeautifulSoup để phân tích cú pháp trang web
        soup = BeautifulSoup(response.content, 'html.parser')
        # Tìm thẻ div có class
        article_title_tag = soup.find('div', class_='article-title')
        article_title_tag_1 = soup.find('h2', class_='articleTitle cABlue')
        if article_title_tag:
            article_title = article_title_tag.text.strip()
        elif article_title_tag_1:
            article_title = article_title_tag_1.text.strip()
        return article_title

def write_en_article_to_doc(url):
    # Gửi yêu cầu HTTP để lấy nội dung trang web
    response = requests.get(url)
    english_text = []
    article_title = "Can not get artical title from web"
    if response.status_code == 200:
        # Sử dụng BeautifulSoup để phân tích cú pháp trang web
        soup = BeautifulSoup(response.content, 'html.parser')
        # Tìm thẻ div có class
        article_contents = soup.find('div', class_='article-body-content')
        if article_contents:
            paragraphs = article_contents.find_all(['p', 'h3'])
        article_title_tag = soup.find('div', class_='article-title')
        article_title_tag_1 = soup.find('h2', class_='articleTitle cABlue')
        if article_title_tag:
            article_title = article_title_tag.text.strip()
        elif article_title_tag_1:
            article_title = article_title_tag_1.text.strip()
        title_name = article_title
        article_title_tag_byline = soup.find('div', class_='article-byline')
        if not article_title_tag_byline:
              article_title_tag_byline = soup.find('div', class_='dateShare cf')
        article_title_line = article_title_tag_byline.text.strip()
        parts = article_title_line.split("|")
        if len(parts) > 1:
            article_title_line = parts[1].strip()
        else:
            article_title_line = article_title_tag_byline.text.strip()

        div_elements = soup.find_all("div", class_="translation-and-category-info")

        # Lặp qua từng thẻ div
        for div in div_elements:
            # Tìm thẻ a với lớp "cBBlue" bên trong thẻ div
            a_element = div.find("a", class_="cBBlue")

            # Kiểm tra nếu thẻ a tồn tại
            if a_element:
                link = a_element["href"]

        # Tạo tập tin Word mới
        # In đậm cho tiêu đề bài báo
        english_text.append(article_title)
        # In nghiêng cho dòng article_title_line
        english_text.append(article_title_line)
        # Thêm các đoạn văn và tiêu đề vào tập tin Word
        for paragraph in paragraphs:
            if paragraph.name == 'h3':
                english_text.append(paragraph.get_text())
            elif 'splitted' in paragraph.get("class", []):
                span_sections = paragraph.find_all('span', class_='section')
                for span in span_sections:
                    img = span.find('img')
                    if img and img.has_attr('src'):
                        image_link = img['src']
                        image_link_final = "https://en.minghui.org" + image_link
                        english_text.append(image_link_final)
                    else:
                        english_text.append(span.get_text())
            else:
                english_text.append(paragraph.get_text())

        english_text.append(copyright_text)
        link_en = "Bản tiếng Anh: " + url
        link_cn = "Bản tiếng Hán: " + link
        english_text.append(link_cn)
        english_text.append(link_en)
        return english_text, link_en, link_cn

    else:
        return "Can get article content"

def tokenize_sentences_with_name_prefix(text):
    name_prefixes = ["Mr.", "Ms.", "No."]

    # Replace name prefixes with placeholders
    for prefix in name_prefixes:
        text = text.replace(prefix, f"{prefix}DOT")

    # Tokenize sentences
    sentences = nltk.sent_tokenize(text)

    # Replace the placeholders back to name prefixes
    for prefix in name_prefixes:
        sentences = [sentence.replace(f"{prefix}DOT", prefix) for sentence in sentences]

    return sentences

def remove_empty_paragraphs(doc):
    # Tạo danh sách mới để lưu các đoạn không trống
    non_empty_paragraphs = []

    # Duyệt qua từng đoạn trong tài liệu
    for paragraph in doc.paragraphs:
        paragraph.text = re.sub(r'^\(Minghui\.org\)', '', paragraph.text)
        # Kiểm tra xem đoạn có nội dung không
        if paragraph.text.strip():  # Kiểm tra xem đoạn có ít nhất một ký tự không phải khoảng trắng hay không
            non_empty_paragraphs.append(paragraph)

    return non_empty_paragraphs

def paragraph_execute_text(english_paragragh):
    english_text = english_paragragh
    english_sentence_list = tokenize_sentences_with_name_prefix(english_text)
    vietnamese_sentence_list =""
    for english_sentence in english_sentence_list:
        vietnamese_sentence = find_vietnamese_sentence(english_sentence)
        if vietnamese_sentence:
            vietnamese_sentence_list = vietnamese_sentence_list + " " +vietnamese_sentence
        else:
            vietnamese_sentence_list = vietnamese_sentence_list + " " + translate_with_google_translate(english_sentence)
    vietnamese_sentence_list = vietnamese_sentence_list.strip()

    return vietnamese_sentence_list

def paragraph_text_check(english_paragragh):
    link = False
    related_text = False
    copyright_check = False
    text_en = english_paragragh
    if 'Related Report' in text_en or 'Related report' in text_en:
        related_text = True
    if "Related Report" in text_en or "Related report" in text_en:
        related_text = True
    if 'en.minghui.org' in text_en or "en.minghui.org" in text_en:
        link = True
    if "Bản quyền" in text_en or 'Bản quyền' in text_en:
        copyright_check = True

    return link, related_text, copyright_check

def write_related_link_to_doc(url, link_en, link_cn):

    vietnamese_text = []
    # Khởi động ứng dụng Microsoft Word
    english_text, vietnam_link, english_link = find_vietnamese_link(url)

    for i in range(len(english_text)):
        vietnamese_text.append(english_text[i])
        vietnamese_text.append(vietnam_link[i])
    vietnamese_text.append(copyright_text)
    vietnamese_text.append(link_cn)
    vietnamese_text.append(link_en)
    return vietnamese_text

"""
def write_related_link_to_doc_1(english_text, vietnam_link, english_link, title):
    vietnamese_text = []
    project_folder = os.getcwd()
    # Khởi động ứng dụng Microsoft Word
    file_name = "Related article links.docx"
    document_path = os.path.join(project_folder, file_name)
    doc = Document()
    doc.add_paragraph(title)
    #english_text, vietnam_link, english_link = find_vietnamese_link(url)
    for link in vietnam_link:
        vietnamese_text.append(get_vn_article_title(link))
    for i in range(len(english_text)):
            doc.add_paragraph(english_text[i])
            doc.add_paragraph(vietnam_link[i])
    doc.save(file_name)
    # Đóng ứng dụng Microsoft Word
    return document_path
"""

def read_paragraph_in_word(url):
    translate_text = []
    english_text = []
    english_article, link_en, link_cn = write_en_article_to_doc(url)
    for item in english_article:
        text = re.sub(r'^\(Minghui\.org\)', '', item)
        # Kiểm tra xem đoạn có nội dung không
        if text:  # Kiểm tra xem đoạn có ít nhất một ký tự không phải khoảng trắng hay không
            english_text.append(text)

    related_article = False
    check_done = False
    # Lặp qua từng đoạn trong tệp văn bản
    for item in english_text:
        link, related_text, copyright_check = paragraph_text_check(item)
        if copyright_check:
            check_done = True
            break
        if related_text:
            related_article = True
            check_done = True
            break
        if link:
            translate_text.append(item)
        else:
            vietnamese_paragraph = paragraph_execute_text(item)
            translate_text.append(item)
            translate_text.append((vietnamese_paragraph))

    if related_article:
        end_text = write_related_link_to_doc(url, link_en, link_cn)
        for item in end_text:
            translate_text.append(item)

    else:

        translate_text.append(copyright_text)
        translate_text.append(link_cn)
        translate_text.append(link_en)

    return translate_text

# Đường dẫn đến bài báo
#url = "https://en.minghui.org/html/articles/2023/7/15/210315.html"

#file_test = "my_document.docx"

#write_related_link_to_doc(file_test, url)
#write_en_article_to_doc(url)
#read_paragraph_in_word(title_name ,url)
#print(get_en_article_title(url))





