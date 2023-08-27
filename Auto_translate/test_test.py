import nltk
import os
import pandas as pd
from docx import Document
import requests
from bs4 import BeautifulSoup
from google_translate import translate_with_google_translate
import win32com.client

title_name = ""
link_en = ""
link_cn = ""
copyright_text = "Bản quyền © 2023 Minghui.org. Mọi quyền được bảo lưu."

def get_related_link(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        article_contents_1 = soup.find('div', class_='article-body-content')
        article_contents_2 = soup.find('div', class_='articleZhengwen geo cBBlack')
        if article_contents_1:
            article_contents = article_contents_1
        else:
            article_contents = article_contents_2
        if article_contents:
            paragraphs = article_contents.find_all(['p', 'h3'])
            valid_paragraphs = []
            start_collecting = False
            related_links = []
            for paragraph in paragraphs:
                text = paragraph.get_text(strip=True)

                if "Related article" in text or "Related Article" in text or "Related Report" in text or "Related report" in text:
                    start_collecting = True

                if start_collecting:
                    if 'splitted' in paragraph.get("class", []):
                        span_sections = paragraph.find_all('span', class_='section')
                        for span in span_sections:
                            text_1 = span.get_text(strip=True)
                            valid_paragraphs.append(text_1)
                            for link in paragraph.find_all('a', href=True):
                                related_links.append(link['href'])
                    else:
                        valid_paragraphs.append(text)
                        for link in paragraph.find_all('a', href=True):
                            related_links.append(link['href'])

            if valid_paragraphs:
                article_content = "\n".join(valid_paragraphs)
                return valid_paragraphs, related_links
            else:
                return [], related_links
        else:
            return [], []
    else:
        return [], []

def find_vietnamese_link(english_link):
    first_text = "Bài liên quan:"
    all_link = []
    result_link = []
    result_link.append(first_text)
    link_fail = "Can not find Vietnamese Link"
    # Đọc dữ liệu từ file CSV
    csv_filename = "link_eng_vn_gct.csv"
    df = pd.read_csv(csv_filename)
    related_content, related_link = get_related_link(english_link)
    for link in related_link:
        all_link.append(link)

    for link in all_link:
        # Tìm link tiếng Anh trong cột 1
        row = df[df.iloc[:, 0] == link]

        if not row.empty:
            vietnamese_link = row.iloc[0, 1]
            result_link.append(vietnamese_link)
        else:
            result_link.append(link_fail)
    return related_content, result_link, all_link

def find_vietnamese_sentence(english_sentence):

    # Đọc dữ liệu từ file CSV
    csv_filename = "dic_eng_vn_data.csv"
    df = pd.read_csv(csv_filename)

    # Tìm link tiếng Anh trong cột 1
    row = df[df.iloc[:, 0] == english_sentence]

    if not row.empty:
        vietnamese_sentence = row.iloc[0, 1]
        return vietnamese_sentence
    else:
        return []
def get_vn_article_title(url):
    try:
        # Tải nội dung của trang web
        response = requests.get(url)

        # Kiểm tra nếu yêu cầu thành công (status code 200)
        if response.status_code == 200:
            # Sử dụng BeautifulSoup để phân tích cú pháp trang web
            soup = BeautifulSoup(response.content, 'html.parser')

            # Tìm thẻ div có class là 'article-title'
            article_title_tag = soup.find('h1', class_='articleTitle cABlue')

            # Kiểm tra nếu tồn tại thẻ và lấy nội dung text của tiêu đề
            if article_title_tag:
                article_title = article_title_tag.text.strip()
                return article_title
            else:
                return ""
        else:
            return ""
    except Exception as e:
        return ""

def write_en_article_to_doc(url):
    global title_name
    global link_en
    global link_cn
    # Gửi yêu cầu HTTP để lấy nội dung trang web
    response = requests.get(url)
    if response.status_code == 200:
        # Sử dụng BeautifulSoup để phân tích cú pháp trang web
        soup = BeautifulSoup(response.content, 'html.parser')
        # Tìm thẻ div có class
        article_contents = soup.find('div', class_='article-body-content')
        if article_contents:
            paragraphs = article_contents.find_all(['p', 'h3'])
        article_title_tag = soup.find('div', class_='article-title')
        if article_title_tag:
            article_title = article_title_tag.text.strip()
        title_name = article_title
        article_title_tag_byline = soup.find('div', class_='article-byline')
        article_title_line = article_title_tag_byline.text.strip()
        parts = article_title_line.split("|")
        if len(parts) > 1:
            article_title_line = parts[1].strip()
        else:
            article_title_line = article_title_tag_byline.text.strip()

        div_elements = soup.find_all("div", class_="translation-and-category-info")

        # Lặp qua từng thẻ div
        for div in div_elements:
            # Tìm thẻ a với lớp "cBBlue" bên trong thẻ div
            a_element = div.find("a", class_="cBBlue")

            # Kiểm tra nếu thẻ a tồn tại
            if a_element:
                link = a_element["href"]

        # Tạo tập tin Word mới
        doc = Document()
        # In đậm cho tiêu đề bài báo
        bold_run = doc.add_paragraph().add_run(article_title)
        bold_run.font.bold = True

        # In nghiêng cho dòng article_title_line
        italic_run = doc.add_paragraph().add_run(article_title_line)
        italic_run.font.italic = True

        # Thêm các đoạn văn và tiêu đề vào tập tin Word
        for paragraph in paragraphs:
            if paragraph.name == 'h3':
                run = doc.add_paragraph().add_run(paragraph.get_text())
                run.font.bold = True
            elif 'splitted' in paragraph.get("class", []):
                span_sections = paragraph.find_all('span', class_='section')
                for span in span_sections:
                    img = span.find('img')
                    if img and img.has_attr('src'):
                        image_link = img['src']
                        image_link_final = "https://en.minghui.org" + image_link
                        doc.add_paragraph(image_link_final)
                    else:
                        doc.add_paragraph(span.get_text())
            else:
                p = doc.add_paragraph()
                p.add_run(paragraph.get_text())



        italic_run = doc.add_paragraph().add_run(copyright_text)
        italic_run.font.italic = True

        doc.add_heading("", 0)
        link_en = "Bản tiếng Anh: " + url
        doc.add_paragraph(link_en)
        link_cn = "Bản tiếng Hán: " + link
        doc.add_paragraph(link_cn)

        doc.save(article_title + '.docx')

    else:
        return None

def tokenize_sentences_with_name_prefix(text):
    name_prefixes = ["Mr.", "Ms.", "No."]

    # Replace name prefixes with placeholders
    for prefix in name_prefixes:
        text = text.replace(prefix, f"{prefix}DOT")

    # Tokenize sentences
    sentences = nltk.sent_tokenize(text)

    # Replace the placeholders back to name prefixes
    for prefix in name_prefixes:
        sentences = [sentence.replace(f"{prefix}DOT", prefix) for sentence in sentences]

    return sentences

def remove_empty_paragraphs(doc):
    # Tạo danh sách mới để lưu các đoạn không trống
    non_empty_paragraphs = []

    # Duyệt qua từng đoạn trong tài liệu
    for paragraph in doc.paragraphs:
        # Kiểm tra xem đoạn có nội dung không
        if paragraph.text.strip():  # Kiểm tra xem đoạn có ít nhất một ký tự không phải khoảng trắng hay không
            non_empty_paragraphs.append(paragraph)

    return non_empty_paragraphs

def paragraph_execute_text(english_paragragh):
    english_text = english_paragragh.text.strip()
    english_sentence_list = tokenize_sentences_with_name_prefix(english_text)
    vietnamese_sentence_list =""
    for english_sentence in english_sentence_list:
        vietnamese_sentence = find_vietnamese_sentence(english_sentence)
        if vietnamese_sentence:
            vietnamese_sentence_list = vietnamese_sentence_list + " " +vietnamese_sentence
        else:
            vietnamese_sentence_list = vietnamese_sentence_list + " " + translate_with_google_translate(english_sentence)

    return  vietnamese_sentence_list

def paragraph_text_check(english_paragragh):
    link = False
    related_text = False
    copyright_check = False
    text_en = english_paragragh.text.strip()
    if 'Related Report' in text_en or 'Related report' in text_en:
        related_text = True
    if "Related Report" in text_en or "Related report" in text_en:
        related_text = True
    if 'en.minghui.org' in text_en or "en.minghui.org" in text_en:
        link = True
    if "Bản quyền" in text_en or 'Bản quyền' in text_en:
        copyright_check = True

    return link, related_text, copyright_check

def write_related_link_to_doc(file_name, url):
    vietnamese_text = []
    project_folder = os.path.dirname(os.path.abspath(__file__))
    # Khởi động ứng dụng Microsoft Word
    word = win32com.client.Dispatch("Word.Application")
    # Mở tài liệu Word từ thư mục của dự án
    document_path = os.path.join(project_folder, file_name)
    doc = word.Documents.Open(document_path)
    # Di chuyển con trỏ đến cuối tài liệu
    doc.Range().Collapse(0)  # Điểm bắt đầu
    # Thêm một dòng trống (enter) để chuyển xuống dòng mới
    doc.Range().InsertAfter('\n')
    english_text, vietnam_link, english_link = find_vietnamese_link(url)
    for link in vietnam_link:
        vietnamese_text.append(get_vn_article_title(link))
    for i in range(len(english_text)):
        if i == 0:
            #write Eng
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = english_text[i]
            doc.Range().InsertAfter('\n')
            #write vn
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = vietnam_link[i]
            doc.Range().InsertAfter('\n')
        else:
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = english_text[i]
            # Thêm hyperlink
            hyperlink = new_paragraph.Range.Hyperlinks.Add(Anchor=new_paragraph.Range,

                                                           Address=english_link[i-1], SubAddress="",
                                                           TextToDisplay="Nhấn vào đây")
            doc.Range().InsertAfter('\n')
            #Vietnam
            new_paragraph = doc.Content.Paragraphs.Add()
            new_paragraph.Range.Text = vietnamese_text[i]
            # Thêm hyperlink
            hyperlink = new_paragraph.Range.Hyperlinks.Add(Anchor=new_paragraph.Range,

                                                           Address=vietnam_link[i], SubAddress="",
                                                           TextToDisplay="Nhấn vào đây")
            doc.Range().InsertAfter('\n')

    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = copyright_text
    doc.Range().InsertAfter('\n')

    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = link_en
    doc.Range().InsertAfter('\n')

    new_paragraph = doc.Content.Paragraphs.Add()
    new_paragraph.Range.Text = link_cn
    doc.Range().InsertAfter('\n')

    doc.Save()
    # Đóng tài liệu
    doc.Close()
    # Đóng ứng dụng Microsoft Word
    word.Quit()

def read_paragraph_in_word(file_name, url):
    file_name_translate = file_name + "_translate"
    related_article = False
    # Đọc tệp văn bản
    doc = Document(file_name + '.docx')
    non_empty = remove_empty_paragraphs(doc)
    doc_translate = Document()
    doc_translate.add_paragraph().add_run("Created by Auto_translate tool").bold = True
    doc_translate.save(file_name_translate + '.docx')
    # Lặp qua từng đoạn trong tệp văn bản
    for paragraph in non_empty:
        link, related_text, copyright_check = paragraph_text_check(paragraph)
        if copyright_check:
            break
        if related_text:
            related_article = True
            break
        doc_translate = Document(file_name_translate + '.docx')
        if link:
            doc_translate.add_paragraph(paragraph.text.strip())
        else:
            vietnamese_paragraph = paragraph_execute_text(paragraph)
            doc_translate.add_paragraph(paragraph.text.strip())
            doc_translate.add_paragraph(vietnamese_paragraph)
        #if i == 3:
           #break
        doc_translate.save(file_name_translate + '.docx')
    doc_translate.save(file_name_translate + '.docx')
    if related_article:
        write_related_link_to_doc(file_name_translate + '.docx', url)

    else:
        doc_translate = Document(file_name_translate + '.docx')

        italic_run = doc.add_paragraph().add_run(copyright_text)
        italic_run.font.italic = True
        doc.add_heading("", 0)
        doc.add_paragraph(link_en)
        doc.add_paragraph(link_cn)
        doc_translate.save(file_name_translate + '.docx')

    return None

# Đường dẫn đến bài báo
url = "https://en.minghui.org/html/articles/2023/7/15/210315.html"

file_test = "my_document.docx"

#write_related_link_to_doc(file_test, url)
write_en_article_to_doc(url)
read_paragraph_in_word(title_name ,url)






