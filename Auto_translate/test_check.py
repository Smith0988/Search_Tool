from docx import Document
from docx.enum.text import WD_ALIGN
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import webbrowser

# Tạo một tệp Word mới
doc = Document()

# Tạo đối tượng hyperlink
hyperlink = doc.add_paragraph().add_run('Visit OpenAI')
hyperlink.hyperlink.address = 'https://www.openai.com/'

# Điều chỉnh kiểu và màu sắc của hyperlink
hyperlink.font.size = Pt(12)
hyperlink.font.color.rgb = Pt(0x00, 0x00, 0xFF)  # Màu xanh

# Cách để hyperlink trung tâm trên dòng
hyperlink._r.get_or_add_rPr().append(
    parse_xml(r'<w:rStyle xmlns:w="%s"><w:rFonts xmlns:w="%s" w:ascii="Times New Roman"/></w:rStyle>' %
              (nsdecls('w'), nsdecls('w')))
)
hyperlink._r.get_or_add_rPr().append(
    parse_xml(r'<w:jc xmlns:w="%s" w:val="center"/>' % nsdecls('w'))
)

# Lưu tệp Word
doc.save('hyperlink.docx')

# Mở tệp Word sau khi đã lưu
webbrowser.open('hyperlink.docx')
